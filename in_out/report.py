from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import tempfile, os
from models import PumpingTest
from runner import ConstantRateSession, RecoverySession, StepDrawdownSession
from analysis.interpretation import interpret_constant_rate, interpret_recovery, interpret_step_drawdown

def generate_report(session) -> BytesIO:
    doc = Document()
        
    # Title
    doc.add_heading("Pumping Test Analysis Report", 0)
    doc.add_paragraph(f"Borehole: {session.test.borehole.name}")
    if session.test.test_date:
        doc.add_paragraph(f"Date: {session.test.test_date}")
    if session.test.operator:
        doc.add_paragraph(f"Operator: {session.test.operator}")
    
    doc.add_heading("Results", level=1)
    
    # Results table
    if isinstance(session, ConstantRateSession):
        r = session.result
        rows = [
            ("Transmissivity", f"{r.transmissivity_m2day:.2f}", "m²/day"),
            ("Estimated Yield", f"{r.estimated_yield_m3day:.2f}", "m³/day"),
            ("Flowrate", f"{r.flowrate_m3day/24:.2f}", "m³/h"),
            ("Drawdown per log cycle", f"{r.fit.drawdown_per_log_cycle:.3f}", "m"),
            ("R²", f"{r.fit.r_squared:.4f}", ""),
        ]
        interpretation = interpret_constant_rate(r, session.test.borehole.name)
    elif isinstance(session, RecoverySession):
        r = session.result
        rows = [
            ("Transmissivity", f"{r.transmissivity_m2day:.2f}", "m²/day"),
            ("Estimated Yield", f"{r.estimated_yield_m3day:.2f}", "m³/day"),
            ("Final Recovery", f"{r.recovery_pcg:.1f}", "%"),
            ("R²", f"{r.fit.r_squared:.4f}", ""),
        ]
        interpretation = interpret_recovery(r, session.test.borehole.name)
    else:
        r = session.result
        rows = [
            ("Aquifer Loss Coefficient (B)", f"{r.aquifer_loss_coeff:.4f}", "m/(m³/h)"),
            ("Well Loss Coefficient (C)", f"{r.well_loss_coeff:.4f}", "m/(m³/h)²"),
            ("Critical Yield", f"{r.critical_yield_m3h:.2f}", "m³/h"),
            ("R²", f"{r.r_squared:.4f}", ""),
        ]
        interpretation = interpret_step_drawdown(r, session.test.borehole.name)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Parameter", "Value", "Units"
    for param, value, unit in rows:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = param, value, unit
    
    doc.add_heading("Interpretation", level=1)
    # Strip markdown bold markers for plain docx paragraph
    doc.add_paragraph(interpretation.replace("**", ""))
    
    doc.add_heading("Reference", level=1)
    doc.add_paragraph(
        "ICRC (2011). Technical Review — Practical Guidelines for Test Pumping in Water Wells. "
        "International Committee of the Red Cross, Geneva."
    )
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf